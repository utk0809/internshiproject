import sys
import subprocess

def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])
        __import__(package)

install_and_import('docx')
from docx import Document
from docx.shared import Pt
import datetime

def create_report(filename, title, points):
    doc = Document()
    doc.add_heading(title, 0)
    
    for point in points:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(point).bold = False
        
    doc.save(filename)
    print(f"Created {filename}")

# Doc 1: Login Task
login_points = [
    "Task Name: Implement Fake Form / Perfect Instagram Clone",
    "Actions Taken: Converted HTML to PHP. Replaced the basic layout with a perfect 2-column dark-mode desktop replica of the live official Instagram login page. This included fetching genuine CDN assets, applying precise CSS gradient masks, fixing reactive floating label overflow issues, and injecting custom JS logic to exactly sync the 5-second rotating phone screenshot animations.",
    "Time Taken: Approximately 3 Hours"
]
create_report(r"d:\CS T02\Login_Task_Report.docx", "Task 1 Report: Phishing Login Page", login_points)

# Doc 2: Awareness Task
awareness_points = [
    "Task Name: Cyber Security Awareness Portal",
    "Actions Taken: Redesigned the placeholder page into a premium, interactive dark-mode cyber security educational dashboard using modern typography (Outfit) and FontAwesome icons. Mapped authentic learning objectives from the project PDF directly into the content, writing original, human-sounding explanations regarding Social Engineering, visual URL Spoofing differentiation, and mandatory User Defenses (like MFA and Password Managers).",
    "Time Taken: Approximately 2 Hours"
]
create_report(r"d:\CS T02\Awareness_Task_Report.docx", "Task 2 Report: Cyber Security Awareness", awareness_points)
